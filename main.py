import os
import time
from PIL import Image
from docx import Document
from docx.shared import Inches


def place_photo_on_template(photo_path, template_path, output_path):
    # Fotoğrafı yükle
    photo = Image.open(photo_path)

    # Şablonu yükle
    doc = Document(template_path)

    # Fotoğrafı boyutlandır
    photo.thumbnail((500, 750))  # Fotoğrafı istenen boyuta boyutlandır

    # Fotoğrafı şablona yerleştir
    table = doc.tables[0]  # Şablonun ilk tablosunu al
    cell = table.cell(0, 0)  # İstenilen hücreyi seç
    cell.text = ""  # Hücre içeriğini temizle

    run = cell.paragraphs[0].add_run()
    run.add_picture(photo_path, width=Inches(5.1), height=Inches(4.5))

    left = Inches(2.0)  # İstenilen sol kenar boşluğu
    top = Inches(1.9)  # İstenilen üst kenar boşluğu
    photo.left = int(left * 9525)
    photo.top = int(top * 9525)

    # Sonucu kaydet
    doc.save(output_path)


def monitor_folder(folder_path, template_path, output_folder):
    processed_files = set()  # İşlenen dosyaları takip etmek için bir küme oluştur

    while True:
        files = os.listdir(folder_path)

        for file in files:
            file_path = os.path.join(folder_path, file)
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(file_path)[1].lower()
                if file_ext in ['.jpg', '.jpeg', '.png']:
                    # Dosya daha önce işlenmediyse işle
                    if file_path not in processed_files:
                        output_path = os.path.join(output_folder, f"output_{file}.docx")
                        place_photo_on_template(file_path, template_path, output_path)
                        processed_files.add(file_path)
                        print(f"Yeni bir fotoğraf işlendi: {file_path}")

        time.sleep(1)


def main():
    folder_path = r"C:\Users\Oğuzhan\PycharmProjects\Fotografbaski\photo"
    template_path = r"C:\Users\Oğuzhan\PycharmProjects\Fotografbaski\template\template1.docx"
    output_folder = "output_folder"

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    monitor_folder(folder_path, template_path, output_folder)


if _name_ == '_main_':
    main()